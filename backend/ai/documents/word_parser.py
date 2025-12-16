from docx import Document
from pathlib import Path
from typing import Dict, Tuple
import logging

logger = logging.getLogger(__name__)

class WordParser:
    """Parse Word (.docx) files"""
    
    @staticmethod
    def parse_word(file_path: str) -> Tuple[str, str]:
        """
        Parse Word file and return (content, format)
        
        Returns:
            (content, format_type)
        """
        file_name = Path(file_path).name
        
        try:
            doc = Document(file_path)
            logger.info(f"Reading {file_name} - {len(doc.paragraphs)} paragraphs")
            
            # Extract all text content
            content_parts = []
            
            # Add paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Skip empty paragraphs
                    content_parts.append(text)
            
            # Add tables
            for table in doc.tables:
                content_parts.append("\n--- Table ---")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    content_parts.append(row_text)
                content_parts.append("--- End Table ---\n")
            
            content = "\n".join(content_parts)
            
            logger.info(f"  ✓ Extracted {len(content)} characters")
            return content, 'text'
            
        except Exception as e:
            logger.error(f"Failed to parse {file_name}: {e}")
            raise
    
    @staticmethod
    def get_file_info(file_path: str) -> Dict:
        """Get file metadata"""
        path = Path(file_path)
        return {
            'file_name': path.name,
            'file_type': path.suffix,
            'file_size': path.stat().st_size,
            'exists': path.exists()
        }